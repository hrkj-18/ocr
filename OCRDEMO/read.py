from PIL import Image                                                                                                                                 
import glob
import os
import cv2
from scipy.misc import imsave, imread, imresize
import sys
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'

#from keras.layers import Conv2D, MaxPooling2D, Convolution2D, Dropout, Dense, Flatten, LSTM
from keras.models import Sequential, save_model, load_model
from keras.utils import np_utils
from scipy.io import loadmat
import pickle
import argparse
import keras
import numpy as np
import argparse
from keras.models import model_from_yaml
import re
import base64
import pickle
import h5py
from docx import Document
from docx.shared import Inches
import grammar_ginger
#from grammar_ginger import main
ap = argparse.ArgumentParser()
ap.add_argument("-tl", "--template_label", required=True, help="Path to the image")
ap.add_argument("-tt", "--template_type", required=True, help="Path to the image")
ap.add_argument("-ta", "--template_autocorrect", required=True, help="Path to the image")
args = vars(ap.parse_args())

template_label = args["template_label"]
with open(template_label) as f:
       content_list=f.read().split(',')
labels=[]
for x in content_list[:-1]:
       labels.append(x) 

print(labels)

template_type = args["template_type"]
with open(template_type) as f:
       content_list=f.read().split(',')
types=[]
for x in content_list[:-1]:
       types.append(int(x)) 

print(types)

template_autocorrect = args["template_autocorrect"]
with open(template_autocorrect) as f:
       content_list=f.read().split(',')
autos=[]
for x in content_list[:-1]:
       autos.append(int(x)) 

print(autos)

def load_model(bin_dir):
    
      '''
      Load model from .yaml and the weights from .h5

          Arguments:
              bin_dir: The directory of the bin (normally bin/)

          Returns:
              Loaded model from file

      '''
      # load YAML and create model
      yaml_file = open('%s/model.yaml' % bin_dir, 'r')
      loaded_model_yaml = yaml_file.read()
      yaml_file.close()
      model = model_from_yaml(loaded_model_yaml)

      # load weights into new model
      model.load_weights('%s/model.h5' % bin_dir)
      return model

bin_dir = r'/home/simran/Desktop/OCRDEMO/bin'
model  = load_model(bin_dir)
bin_dir = r'/home/simran/Desktop/OCRDEMO/char_bin'
char_model = load_model(bin_dir)
bin_dir = r'/home/simran/Desktop/OCRDEMO/num_bin'
num_model = load_model(bin_dir)
#model = load_model('/home/simran/Keras_Alphabet_Model.h5')


dic1 = {0:'0',
       1:'1',
       2:'2',
       3:'3',
       4:'4',
       5:'5',
       6:'6',
       7:'7',
       8:'8',
       9:'9',
       10:'A',
       11:'B',
       12:'C',
       13:'D',
       14:'E',
       15:'F',
       16:'G',
       17:'H',
       18:'I',
       19:'J',
       20:'K',
       21:'L',
       22:'M',
       23:'N',
       24:'O',
       25:'P',
       26:'Q',
       27:'R',
       28:'S',
       29:'T',
       30:'U',
       31:'V',
       32:'W',
       33:'X',
       34:'Y',
       35:'Z',
       36:'a',
       37:'b',
       38:'c',
       39:'d',
       40:'e',
       41:'f',
       42:'g',
       43:'h',
       44:'i',
       45:'j',
       46:'k',
       47:'l',
       48:'m',
       49:'n',
       50:'o',
       51:'p',
       52:'q',
       53:'r',
       54:'s',
       55:'t',
       56:'u',
       57:'v',
       58:'w',
       59:'x',
       60:'y',
       61:'z'
      }
'''
         
dic1 = {0:'A',
       1:'B',
       2:'C',
       3:'D',
       4:'E',
       5:'F',
       6:'G',
       7:'H',
       8:'I',
       9:'J',
       10:'K',
       11:'L',
       12:'M',
       13:'N',
       14:'O',
       15:'P',
       16:'Q',
       17:'R',
       18:'S',
       19:'T',
       20:'U',
       21:'V',
       22:'W',
       23:'X',
       24:'Y',
       25:'Z'
       }
'''

dtypeList=[]
dtypeList=types
typeList=[]
for i in range(len(dtypeList)):
      typeList.append(dtypeList[i])
#typeList = types
labelsList=labels
autosList = autos
removed=0
for i in range(len(dtypeList)):
      print(i)
      print(dtypeList)
      if (dtypeList[i]==2):
            
            #print(removed)
            del labelsList[i-removed]
            #print(labelsList)
            del typeList[i-removed]
            print(typeList)
            del autosList[i-removed]
            #print(autosList)
            print(dtypeList)

            removed=removed+1
            

imageFolderPath = '/home/simran/Desktop/OCRDEMO/Images'
n=os.popen('find Images -type d | wc -l').read()
n = int(n)
#print(n)
n=n-1
Lines = []
print(n)
for i in range(0,n):
      imagePath = sorted(glob.glob(imageFolderPath +'/'+str(i).zfill(4)+ '/*.jpg')) 
      line = np.array( [np.array(Image.open(img).convert('L'), 'f') for img in imagePath] )
      #print(i)
      if(np.shape(line)[0]!=0):
            #print(line)
            Lines.append(line)
#print(Lines)
for i in range (len(Lines)):
      Lines[i]=np.reshape(Lines[i],(-1,1,28,28,1))
      #Lines[i]=np.reshape(Lines[i],(-1,1,784))
      Lines[i]=Lines[i].astype('float32')
      Lines[i] /= 255
      #print(np.shape(Lines[i]))


out_string=[[]]
h=0
output_line=[]
output_max=[]
for i in range(len(Lines)):
      for j in range(len(Lines[i])):
            #Lines[i][j].reshape(1,28,28,1)
            #print(np.shape(Lines[i][j]))
            #h=h+1
            #print(h)
            #output=model.predict_proba(Lines[i][j])
            #output=model.predict(Lines[i])
            print("i : "+str(i))
            #print(ddtypeList[i])
            if(int(typeList[i])==1):
                  output=char_model.predict_classes(Lines[i][j])
                  output=output[0]
                  output+=10
                  output_line.append(output)
                  #print(output)
            elif(int(typeList[i])==0):
                  output=num_model.predict_classes(Lines[i][j])
                  output=output[0]
                  output_line.append(output)
            #else:
                 #continue
            #output=model.predict_classes(Lines[i][j])
            #output_line.append(output)
      output_max.append(output_line)
      #output_line[k]=[dic.get(n, n) for n in output_line]
      #print(output_line)
      output_line=[]


print(output_max)
'''
goat=[]
goat_real=[]
for i in range(len(output_max)):
      for j in range(len(output_max[i])):
            goat.append(output_max[i][j][0])
      goat_real.append(goat)
      goat=[]
'''
#print(goat_real)
print(output_max)
goat_real=output_max

print(goat_real)

#removed=0
#for i in range(len(dtypeList)):
      #if (dtypeList[i]==2):
            #del goat_real[i-removed]
            #removed=removed+1
print(goat_real)

new_goat=[]
for i in range(len(goat_real)):
    new_goat.append([dic1.get(n, n) for n in goat_real[i]])

print(new_goat)

for i in range(len(new_goat)):
    for j in range(len(new_goat[i])):
        new_goat[i]=''.join(new_goat[i])
print(new_goat)

print(labelsList)
print(dtypeList)
print(autosList)

for i in range(len(new_goat)):
      if(autosList[i]==1):
            lower=new_goat[i].lower()
            print(new_goat[i])
            new_goat[i]=grammar_ginger.main(lower)
            new_goat[i] = ''.join(e for e in new_goat[i] if e.isalnum())
            #mpa = dict.fromkeys(range(32))
            #neww[i]=neww[i].translate(mpa)
            new_goat[i]=new_goat[i][:-2]
            new_goat[i]=new_goat[i].upper()
            print(new_goat[i])
      else:
            new_goat[i]=new_goat[i]
neww=[]
for i in new_goat:
      neww.append(i)

#print(new_goat)
'''
for i in range(len(new_goat)):
      print(new_goat[i])
'''
#print("##################################################################")
newl=labelsList
print(newl)
document = Document()
p = document.add_heading('TEST FORM')
#for i in range(len(newl)):
      #print(newl[i]+' : '+neww[i])

dtype=dtypeList

#for i in range(len(newl)):
#      document.add_paragraph(newl[i]+' : '+neww[i])

#document.add_picture('D_images/1.jpg', width=Inches(1.25))
j=0
imread=0
for i in range(len(dtype)):
      if(dtype[i]==2):
            document.add_picture('D_images/'+str(imread)+'.jpg', width=Inches(1.25))
            imread+=1
      else:
            document.add_paragraph(newl[j]+' : '+neww[j])
            print(newl[j]+' : '+neww[j])
            j=j+1


document.add_page_break()
document.save('demo.docx')


orig_stdout = sys.stdout
f = open('/home/simran/Desktop/OCRDEMO/out.txt', 'w')
sys.stdout = f
i=0


for i in range(len(new_goat)):
      print(new_goat[i])
      print("\n")
sys.stdout = orig_stdout
f.close()


#Lines.append(line)
#print(Lines)
